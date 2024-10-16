from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_row_to_table(table, key, value):
    row_cells = table.add_row().cells
    row_cells[0].text = key
    cell = row_cells[1]
    cell.text = value
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def generate_business_canvas_model(filename, canvas_model_data):
    doc = Document()

    # Add title
    doc.add_heading('Business Canvas Model', level=1)

    # Add table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.autofit = True

    # Add main column headers
    table.cell(0, 0).text = 'Business Aspect'
    table.cell(0, 1).text = 'Description'

    for section, data in canvas_model_data.items():
        doc.add_heading(section, level=2)
        for key, value in data.items():
            add_row_to_table(table, key, value)

    # Adjust table column widths
    for col in table.columns:
        for cell in col.cells:
            cell.width = Pt(300)

    # Save the document
    doc.save(filename)

if __name__ == "__main__":
    canvas_model_data = {
        "Customer Segments": {
            "Individuals": "Purchase date fruits for personal consumption.",
            "Retailers": "Sellers of date fruits wanting efficient inventory classification."
        },
        "Value Propositions": {
            "Accurate Date Identification": "Precise classification using advanced CNNs.",
            "User-Friendly Mobile App": "Accessible app (DatePal) for capturing and submitting images.",
            "Real-time Classification": "Instant results for an enhanced user experience.",
            "Multiclass Classification": "Identification of various date varieties."
        },
        "Channels": {
            "Mobile App Stores": "Distribution through Android app stores.",
            "Online Platforms": "Utilizing online platforms and social media for promotion."
        },
        "Customer Relationships": {
            "User Education": "Providing in-app guidance for effective use.",
            "Community Engagement": "Establishing a community forum or social media presence."
        },
        "Revenue Streams": {
            "Freemium Model": "Basic image recognition for free, premium features via subscription.",
            "Partnerships": "Exploring partnerships with retailers or date suppliers."
        },
        "Key Resources": {
            "Image Recognition System": "Development and maintenance of CNN-based system.",
            "Mobile App Development Team": "Skilled professionals for DatePal app.",
            "Server Infrastructure": "Resources for handling image data and real-time classification."
        },
        "Key Activities": {
            "Continuous Improvement": "Regularly updating the image recognition model.",
            "Marketing": "Promoting the app through online channels and social media."
        },
        "Key Partnerships": {
            "Date Suppliers/Retailers": "Collaborating for integration and mutual promotion.",
            "Online Platforms": "Partnering to increase visibility and distribution."
        },
        "Cost Structure": {
            "Development Costs": "Investment in initial development.",
            "Marketing Expenses": "Budget for online promotion and potential partnerships.",
            "Maintenance Costs": "Ongoing expenses for server maintenance and user support."
        },
        "User-Friendly Interface": {
            "Design Team": "Hiring or collaborating with a design team for an intuitive UI."
        }
    }

    generate_business_canvas_model("business_canvas_model.docx", canvas_model_data)
