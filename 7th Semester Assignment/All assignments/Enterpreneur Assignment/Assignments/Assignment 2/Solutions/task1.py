from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt

# Create a new Word document
doc = Document()

# Add a title to the document
doc.add_heading('Entrepreneur vs Intrapreneur Comparison', 0)

# Add a table to the document
table = doc.add_table(rows=6, cols=3)
table.allow_autofit = False

# Define cell styles
for row in table.rows:
    for cell in row.cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

# Set cell widths
widths = (2.5, 3.5, 3.5)
for i, width in enumerate(widths):
    cell = table.cell(0, i)
    cell.width = Pt(width)

# Add header cells
header_cells = table.rows[0].cells
header_cells[0].text = 'Aspect'
header_cells[1].text = 'Entrepreneur'
header_cells[2].text = 'Intrapreneur'

# Data for the table
data = [
    ('Ownership', 'Entrepreneurs establish and own their own businesses, which can range from small startups to large corporations. They have the freedom to shape the company\'s vision, mission, and culture.', 'Intrapreneurs are employees within an established organization, and they do not have ownership rights in the company. They work under the authority and direction of the organization\'s leadership.'),
    ('Risk', 'Entrepreneurs typically assume significant financial and personal risks. They invest their own capital and may need to secure loans or funding from external sources. If their business fails, they can suffer financial losses and potentially even personal bankruptcy.', 'Intrapreneurs face lower personal financial risk. While their projects may carry risks, such as failure or potential downsizing, they do not bear the same level of financial liability as entrepreneurs. They have job security within the organization.'),
    ('Control', 'Entrepreneurs have complete control over their business, which includes decision-making, strategy formulation, and operational management. They determine the company\'s direction and have the autonomy to adapt to changing circumstances.', 'Intrapreneurs work within the framework of an existing organization. They must adhere to company policies, procedures, and the overall corporate structure. Their autonomy is limited to their specific projects and initiatives.'),
    ('Funding', 'Entrepreneurs are responsible for securing their own funding, which can come from personal savings, loans, venture capital, angel investors, or crowdfunding. They need to convince external investors to fund their vision.', 'Intrapreneurs typically have access to the resources and funding of their parent organization. They can use the company\'s existing financial and logistical support to implement and develop their projects. External funding is not typically required.'),
    ('Innovation Focus', 'Entrepreneurs focus on creating new businesses and disrupting existing markets with innovative products or services. They are driven by their own vision and passion to build something from the ground up.', 'Intrapreneurs concentrate on fostering innovation and driving change within their current workplace. They work on improving existing processes, products, or services to contribute to the growth and competitiveness of the company they work for. They align their efforts with the organization\'s goals and strategies.')
]

# Add data to the table
for i, (aspect, entrepreneur_text, intrapreneur_text) in enumerate(data):
    cell = table.cell(i + 1, 0)
    cell.text = aspect

    cell = table.cell(i + 1, 1)
    cell.text = entrepreneur_text

    cell = table.cell(i + 1, 2)
    cell.text = intrapreneur_text

    # Adjust the row height to fit the text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(10)
    cell.height = Pt(max(len(entrepreneur_text.split('\n')), len(intrapreneur_text.split('\n'))) * 10)

# Save the document
doc.save('Entrepreneur_vs_Intrapreneur_Comparison.docx')
